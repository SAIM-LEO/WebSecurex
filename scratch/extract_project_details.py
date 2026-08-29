import docx
import os

docx_path = r"e:\WebSecureX.(2)\My project.docx"
output_path = r"e:\WebSecureX.(2)\scratch\project_details.txt"

os.makedirs(os.path.dirname(output_path), exist_ok=True)

try:
    doc = docx.Document(docx_path)
    with open(output_path, "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    f.write(cell.text + " | ")
                f.write("\n")
    print("Extraction successful")
except Exception as e:
    print(f"Error: {e}")
